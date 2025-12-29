from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def create_contract_en():
    doc = Document()
    
    # Style settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'  # English standard font
    font.size = Pt(10.5)

    # Title
    p = doc.add_paragraph('Exclusive Distribution Agreement (DRAFT)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph('This Agreement is made and entered into by and between the parties concerned based on the spirit of equality and mutual benefit.')

    # Article 1
    doc.add_heading('Article 1  Parties', level=1)
    doc.add_paragraph('Party A (Manufacturer/Supplier)\nCompany Name: Travelbird Technology Co. Ltd\nAddress: ROOM 09, 27/F, HO KING COMMERCIAL CENTRE, 2-16 FA YUEN STREET, MONGKOK, KOWLOON, HONG KONG\nPhone: +86 13902321781\nWebsite: https://travelbird.cc/')
    doc.add_paragraph('Party B (Exclusive Distributor in Japan)\nCompany Name: Select Bargain Store\nAddress: 237-2 Takehara, Nagaizumi-cho, Sunto-gun, Shizuoka, 411-0944, Japan\nPhone: +81 55-986-3601 / Mobile: +81 80-3626-0792\nEmail: info@selectbargainstore.com\nURL: https://selectbargainstore.com')

    # Article 2
    doc.add_heading('Article 2  Appointment', level=1)
    doc.add_paragraph('Party A appoints Party B as its exclusive distributor in the territory described in Article 4 for the retail and wholesale of the Products, and for the promotion of the Brand Name, and Party B accepts such appointment.')

    # Article 3
    doc.add_heading('Article 3  Products', level=1)
    doc.add_paragraph('Product Name: Double-Pole Waterproof Tarp and Sunshade\nModel/Specifications:\nMaterial: 210T Polyester, PU2000 Waterproof Coating + Silver Sunshade Coating\nSize: 435 x 285 cm\nAccessories: Ground Pegs, Windproof Ropes, Elastic Ropes, Storage Bag, 2 Gold-plated aluminum alloy support poles, etc.\n*Note: Black trekking poles/stands are not included in the Product.\nOptional: Nomadnest Hammock\nAnd related accessories and products.')

    # Article 4
    doc.add_heading('Article 4  Territory', level=1)
    doc.add_paragraph('Party B shall sell the Products supplied by Party A under this Agreement for the market within Japan, and shall not sell to markets outside of Japan.')

    # Article 5
    doc.add_heading('Article 5  Basic Terms', level=1)
    doc.add_paragraph('5.1 Party A grants Party B the exclusive distribution rights for the Products in Japan. Party A requires Party B to maintain friendly relations with Party A\'s existing business partners.\n5.2 During the term of this Agreement, Party A shall not sell the Products directly or through any third party to consumers or distributors in Japan without the prior written consent of Party B.\n5.3 Party B is responsible for brand promotion and after-sales service, and is prohibited from engaging in any conduct that damages the brand image.\n5.4 Party A shall provide materials, parts, sales materials, technical information, etc., to Party B as necessary.')

    # Article 6
    doc.add_heading('Article 6  Sales Promotion', level=1)
    doc.add_paragraph('6.1 Party A shall provide information on new products, priority access to special prices, and promotional materials (photos, videos, etc.).\n6.2 Party B is responsible for sales promotion through crowdfunding projects (such as MAKUAKE).')

    # Article 7
    doc.add_heading('Article 7  Payment Terms', level=1)
    doc.add_paragraph('7.1 Payment Method: Telegraphic Transfer (T/T) to the bank account designated by Party A.\n7.2 Payment Schedule:\nAdvance Payment: 30% (upon order confirmation)\nBalance: 70% (before shipment)')

    # Article 8
    doc.add_heading('Article 8  MOQ and Pricing', level=1)
    doc.add_paragraph('8.1 The Minimum Order Quantity (MOQ) shall be the number of backers in Japan during the crowdfunding campaign (e.g., MAKUAKE).\n8.2 MOQ for regular wholesale transactions shall be determined by separate consultation.\n8.3 Special conditions for NON-MOQ: Party A shall produce and ship even if the quantity does not meet the MOQ.\n8.4 Price Terms: The transaction price terms under this Agreement shall be EXW (Ex Works).\n8.5 Shipping Costs: Based on EXW terms, all transportation costs and risks from the factory shall be borne by Party B. Details to be discussed per case.')

    # Article 9
    doc.add_heading('Article 9  Delivery', level=1)
    doc.add_paragraph('9.1 Party A shall deliver to the place designated by Party B on the date designated by Party B.\n9.2 Standard Lead Time: Ships within 3-5 business days if in stock; approx. 60 days if out of stock/new production required (to be confirmed per order).\n9.3 In the event of force majeure, the parties shall consult with each other.\n9.4 Handling of Delivery Delays:\nMinor Delay: Acceptance of split delivery, no cancellation right.\nSerious Delay exceeding 30 days: Party B has the right to cancel and claim a refund.')

    # Article 10
    doc.add_heading('Article 10  Quality Assurance and Defective Products', level=1)
    doc.add_paragraph('10.1 If the initial defect rate exceeds 1% per shipment lot, Party A shall consult with Party B and take measures such as free replacement, additional shipment, or discount.\n10.2 If Party B discovers defective products, it shall report to Party A with photos, videos, etc.\n10.3 Return shipping costs for defective products shall be borne by Party A, and any amount advanced by Party B shall be offset.')

    # Article 11
    doc.add_heading('Article 11  Industrial Property Rights', level=1)
    doc.add_paragraph('11.1 During the term of the Agreement, Party B may use Party A\'s trademarks, logos, designs, etc., within the scope of commercial activities.\n11.2 All intellectual property rights such as trademarks belong to Party A.\n11.3 If Party B receives an infringement warning from a third party, it shall promptly notify Party A.')

    # Article 12
    doc.add_heading('Article 12  Advertising and Promotion', level=1)
    doc.add_paragraph('12.1 During the term of the Agreement, Party B shall bear the advertising and promotion costs in Japan.\n12.2 Production of promotional materials (videos, images, text) requires Party A\'s prior approval.')

    # Article 13
    doc.add_heading('Article 13  Term', level=1)
    doc.add_paragraph('13.1 This Agreement shall be effective for 8 months from the date of signature.\n13.2 Unless written notice is given 30 days prior to the expiration of the term, it shall be automatically renewed (in 8-month units).')

    # Article 14
    doc.add_heading('Article 14  Termination', level=1)
    doc.add_paragraph('14.1 In the event of a material breach of contract, the contract may be terminated after a cure period of 30 days following notice.\n14.2 Inventory disposal upon termination shall be determined by mutual consultation.')

    # Article 15
    doc.add_heading('Article 15  Force Majeure', level=1)
    doc.add_paragraph('15.1 No liability shall be assumed for non-performance due to force majeure such as natural disasters, war, epidemics, terrorism, etc.\n15.2 In the event of force majeure, notice and submission of a certificate shall be made within 15 days.\n15.3 If the force majeure continues for more than 3 months, the contract shall be automatically terminated.')

    # Article 16
    doc.add_heading('Article 16  Dispute Resolution', level=1)
    doc.add_paragraph('16.1 Disputes regarding this Agreement shall first be resolved through amicable consultation.\n16.2 If not resolved by consultation, the matter shall be referred to the arbitration institution of the respondent\'s country of residence. In the case of Japan, arbitration shall be filed with the Japan Commercial Arbitration Association (JCAA).\n16.3 The place of arbitration shall be Tokyo, and the language of arbitration shall be Japanese and English.')

    # Article 17
    doc.add_heading('Article 17  Governing Law', level=1)
    doc.add_paragraph('This Agreement shall be governed by the laws of Japan, and its interpretation and performance shall also be subject to the laws of Japan.')

    # Article 18
    doc.add_heading('Article 18  Miscellaneous', level=1)
    doc.add_paragraph('18.1 Amendments, changes, or additions to this Agreement shall be made by written agreement.\n18.2 Matters not stated in this Agreement shall be determined upon separate consultation.\n18.3 This Agreement shall be treated as confidential and shall not be disclosed to anyone other than necessary parties.')

    # Signatures
    doc.add_paragraph('\n')
    doc.add_paragraph('By signing this Agreement, both parties agree to all terms of this Agreement and promise to fulfill them in good faith.')
    doc.add_paragraph('\n')
    
    # Signature Block A
    doc.add_paragraph('Party A (Travelbird Technology Co. Ltd)')
    doc.add_paragraph('Signatory Name: _____________________________')
    doc.add_paragraph('Title: _____________________________')
    doc.add_paragraph('Date: _______________, 2025')
    doc.add_paragraph('Signature (or Seal): _____________________________')
    doc.add_paragraph('\n')

    # Signature Block B
    doc.add_paragraph('Party B (Select Bargain Store)')
    doc.add_paragraph('Signatory Name: Tsutomu Shimizu')
    doc.add_paragraph('Title: Owner')
    doc.add_paragraph('Date: _______________, 2025')
    doc.add_paragraph('Signature (or Seal): _____________________________')

    file_path = 'Exclusive_Distribution_Agreement_DoublePoleTarp_EN_DRAFT B.docx'
    doc.save(file_path)
    print(f"Successfully created {file_path}")

if __name__ == "__main__":
    create_contract_en()
