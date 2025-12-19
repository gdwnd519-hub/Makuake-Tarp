from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def create_contract():
    doc = Document()
    
    # Style settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'MS Mincho' # Or suitable Japanese font
    font.size = Pt(10.5)

    # Title
    p = doc.add_paragraph('独占販売契約書（案）')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph('本契約書は、当事者間で以下の条項に合意し、平等の精神および相互利益に基づき締結される。')

    # Article 1
    doc.add_heading('第1条　当事者', level=1)
    doc.add_paragraph('Party A（製造元・供給元）\n会社名：Travelbird Technology Co. Ltd\n住所：ROOM 09, 27/F, HO KING COMMERCIAL CENTRE, 2-16 FA YUEN STREET, MONGKOK, KOWLOON, HONG KONG\n電話番号：+86 13902321781\n企業サイト：https://travelbird.cc/')
    doc.add_paragraph('Party B（日本国内独占販売代理店）\n会社名：セレクトバーゲンストア\n住所：〒411-0944 静岡県駿東郡長泉町竹原237-2\n電話番号：+81 55-986-3601／携帯：+81 80-3626-0792\nメール：info@selectbargainstore.com\nURL：https://selectbargainstore.com')

    # Article 2
    doc.add_heading('第2条　委任事項', level=1)
    doc.add_paragraph('Aは、Bを第4条記載の領域における総代理店として、小売販売および卸売の受注、ブランドネームの促進を委任し、Bはこれを承諾する。')

    # Article 3
    doc.add_heading('第3条　権限付与対象製品', level=1)
    doc.add_paragraph('製品名：Double-Pole Waterproof Tarp and Sunshade（ダブルポール防水タープ＆サンシェード）\n型番・仕様：\n素材：210T ポリエステル、PU2000 防水コーティング + シルバーサンシェードコーティング\nサイズ：435×285cm\n付属品：アルミ合金製ポール（タープ用）、グランドペグ、防風ロープ、伸縮ロープ、収納バッグ等を含む一式\n※トレッキングポール（スタンド）は製品に含まれない。\nオプション：Nomadnest Hammock（ハンモック）\nおよびこれに関連する付属品・関連商品を対象とする。')

    # Article 4
    doc.add_heading('第4条　権限付与領域', level=1)
    doc.add_paragraph('Bは、本契約に基づきAから供給を受けた本製品を日本国内市場向けに販売するものとし、日本国外市場向けの販売は行わない。')

    # Article 5
    doc.add_heading('第5条　基本契約条項', level=1)
    doc.add_paragraph('5.1　AはBに日本国内における当該製品の総代理店の販売権を与える。Aの既存取引企業とも友好関係を保つことを義務付ける。\n5.2　Aは契約期間中、Bの事前書面同意なく、日本国内の消費者または販売業者に対して本製品を直接または第三者を通じて販売しない。\n5.3　Bはブランドプロモーションおよびアフターサービスの責任を有し、ブランドイメージを損なう行為を禁止する。\n5.4　Aは必要に応じBに材料・部品・販売資料・技術情報等を提供する。')

    # Article 6
    doc.add_heading('第6条　販売促進', level=1)
    doc.add_paragraph('6.1　Aは新商品情報の提供、特別価格への優先アクセス権、販促物・写真・動画等の提供を行う。\n6.2　Bはクラウドファンディングプロジェクト（MAKUAKE等）での販売促進に責任を有する。')

    # Article 7
    doc.add_heading('第7条　支払い条件', level=1)
    doc.add_paragraph('7.1　支払い方法：A指定の銀行口座への電信送金（T/T）\n7.2　支払い期限：\n前払い：30％（注文確認時）\n残金：70％（船積み前）')

    # Article 8
    doc.add_heading('第8条　MOQおよび価格条件', level=1)
    doc.add_paragraph('8.1　クラウドファンディングキャンペーン（MAKUAKE等）における日本国内支援者数を最低注文数量（MOQ）とする。\n8.2　通常卸売取引におけるMOQは別途協議により決定する。\n8.3　NON-MOQ対応時の特別条件：AはMOQ未達の数量でも生産・出荷に対応する。\n8.4　価格条件：本契約における取引価格条件はEXW（工場渡し）とする。\n8.5　送料について：EXW条件に基づき、工場からのすべての輸送費用およびリスクは当事者Bが負担する。詳細は案件ごとに協議する。')

    # Article 9
    doc.add_heading('第9条　納期および配送', level=1)
    doc.add_paragraph('9.1　AはBが指定する期日に指定場所へ納品する。\n9.2　納期の目安：在庫がある場合は3-5営業日以内に発送、在庫切れで新規生産が必要な場合は約60日以内に発送する（詳細は都度確認）。\n9.3　不可抗力発生時は、両者において協議する。\n9.4　納品遅延に関する取扱い：\n軽微な遅延：分割納品を受け入れ、キャンセル権なし。\n30日を超える重大な遅延：Bにキャンセル・返金請求権を認める。')

    # Article 10
    doc.add_heading('第10条　品質保証および不良品対応', level=1)
    doc.add_paragraph('10.1　初期不良率が1出荷ロットあたり1％を超えた場合、AはBと協議のうえ、無償交換・追加出荷・値引き等の措置を講じる。\n10.2　Bは不良品を発見した場合、写真・動画等によりAへ報告する。\n10.3　不良品返送送料はAが負担し、B立替分は相殺処理する。')

    # Article 11
    doc.add_heading('第11条　工業所有権', level=1)
    doc.add_paragraph('11.1　契約期間中、BはAの商標・ロゴ・意匠等を商業活動範囲内で使用できる。\n11.2　商標等の知的財産権はすべてAに帰属する。\n11.3　Bが第三者から侵害警告を受けた場合、速やかにAに通知する。')

    # Article 12
    doc.add_heading('第12条　広告・販促活動', level=1)
    doc.add_paragraph('12.1　契約期間中、Bは日本国内における広告販促費用を負担する。\n12.2　販促用素材（動画・画像・テキスト）の制作はAの事前承認を得ること。')

    # Article 13
    doc.add_heading('第13条　契約有効期間', level=1)
    doc.add_paragraph('13.1　本契約は署名日から6ヶ月間有効とする。\n13.2　有効期間満了30日前までに書面通知がない場合、自動更新（6ヶ月単位）。')

    # Article 14
    doc.add_heading('第14条　契約の終了', level=1)
    doc.add_paragraph('14.1　重大な契約違反があった場合、通知後30日の是正期間経過により契約終了可能。\n14.2　契約終了時の在庫処理は双方協議により決定する。')

    # Article 15
    doc.add_heading('第15条　不可抗力条項', level=1)
    doc.add_paragraph('15.1　天災、戦争、疫病、テロ等の不可抗力により履行不能の場合、責任を負わない。\n15.2　不可抗力発生時は15日以内に通知および証明書提出を行う。\n15.3　不可抗力が3ヶ月以上継続する場合、自動的に契約は解除される。')

    # Article 16
    doc.add_heading('第16条　紛争解決', level=1)
    doc.add_paragraph('16.1　本契約に関する紛争は、まず友好的協議により解決を図る。\n16.2　協議により解決できない場合は、被申立人の居住国の仲裁機関に付託。日本側の場合は日本商事仲裁協会（JCAA）に仲裁を申し立てる。\n16.3　仲裁地は東京、仲裁言語は日本語および英語とする。')

    # Article 17
    doc.add_heading('第17条　準拠法', level=1)
    doc.add_paragraph('本契約は日本法に準拠し、その解釈および履行も日本法によるものとする。')

    # Article 18
    doc.add_heading('第18条　その他', level=1)
    doc.add_paragraph('18.1　本契約の修正・変更・追加は書面合意による。\n18.2　本契約に記載のない事項は別途協議のうえ決定する。\n18.3　本契約は秘密扱いとし、必要な関係者以外には開示しない。')

    # Signatures
    doc.add_paragraph('\n')
    doc.add_paragraph('本契約書に署名することで、両当事者は本契約のすべての条項に同意し、誠実に履行することを約束する。')
    doc.add_paragraph('\n')
    
    # Signature Block A
    doc.add_paragraph('Party A（Travelbird Technology Co. Ltd）')
    doc.add_paragraph('署名者名：_____________________________')
    doc.add_paragraph('役職：_____________________________')
    doc.add_paragraph('署名日：2025年　　月　　日')
    doc.add_paragraph('署名（または押印）：_____________________________')
    doc.add_paragraph('\n')

    # Signature Block B
    doc.add_paragraph('Party B（セレクトバーゲンストア）')
    doc.add_paragraph('署名者名：Tsutomu Shimizu')
    doc.add_paragraph('役職：Owner')
    doc.add_paragraph('署名日：2025年　　月　　日')
    doc.add_paragraph('署名（または押印）：_____________________________')

    file_path = 'Exclusive_Distribution_Agreement_DoublePoleTarp_DRAFT.docx'
    doc.save(file_path)
    print(f"Successfully created {file_path}")

if __name__ == "__main__":
    create_contract()
