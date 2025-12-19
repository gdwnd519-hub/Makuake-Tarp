from docx import Document
import sys

def read_docx(file_path):
    try:
        doc = Document(file_path)
        print(f"--- Content of {file_path} ---")
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        print("--- End of Content ---")
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    file_path = r"c:\Users\clean\OneDrive\ドキュメント\ウェブページ LP\MAKUAKE タープ\Exclusive_Distribution_Agreement_Final.docx"
    output_path = r"c:\Users\clean\OneDrive\ドキュメント\ウェブページ LP\MAKUAKE タープ\final_content.txt"
    
    try:
        doc = Document(file_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                if para.text.strip():
                    f.write(para.text + "\n")
        print(f"Written to {output_path}")
    except Exception as e:
        print(f"Error: {e}")
